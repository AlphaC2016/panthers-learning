# word api.py
#
# Programmer  : Guy, Shaked
# Date        : 03-04/05/2022
#
# word api
# ------------------------------------------------------------------------------------------------------

# import
import docx
import oci
from oci.config import validate_config
from deep_translator import GoogleTranslator
import json
from googleapiclient.discovery import build
from googlesearch import search
import urllib
from bs4 import BeautifulSoup

FILENAME = r"C:\Users\kelle\Desktop\TESTW.docx"
CONFIG_FILE = r"C:\Users\kelle\.oci\config"
USER = "ocid1.user.oc1..aaaaaaaa7mxia3xroif3st7xwoglkyt3cu24h3zohcqpvpc4oo6lrpnmvuzq"
FINGERPRINT = "a7:e6:99:17:39:c9:e4:17:bb:42:f2:81:1c:d7:84:af"
TENANCY = "ocid1.tenancy.oc1..aaaaaaaa7l5viilzwmlk6gmxfyjiektpyusskdbvlwi6bzaqlvxsxy7q2ryq"
REGION = "il-jerusalem-1"
KEY_FILE = r"C:\Users\kelle\Desktop\oracel_keys\kellerg2002@gmail.com_2023-08-23T11_46_39.769Z.pem"
TRUSTED_METADATA = "trusted_metadata.txt"


def read_docx():
    # Reads the docx file and returns it
    print("enter your word directory:\n")
    file = FILENAME
    doc = docx.Document(f"{file}")
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return "\n".join(fulltext)


def create_config():
    # Create a default config using DEFAULT profile in default location
    # config = oci.config.from_file(f"{CONFIG_FILE}")
    config = {
        "user": f"{USER}",
        "fingerprint": f"{FINGERPRINT}",
        "tenancy": f"{TENANCY}",
        "region": f"{REGION}",
        "key_file": f"{KEY_FILE}"
    }
    validate_config(config)
    return config


def api_ai(translated, config):
    # Get the main key phrases
    ai_language_client = oci.ai_language.AIServiceLanguageClient(config)

    batch_detect_language_key_phrases_response = ai_language_client.batch_detect_language_key_phrases(
        batch_detect_language_key_phrases_details=oci.ai_language.models.BatchDetectLanguageKeyPhrasesDetails(
            documents=[
                oci.ai_language.models.TextDocument(
                    key="doc1",
                    text=f"{translated}",
                    language_code="en")]
        ))
    # JSON that contains a dict with all the main key phrases and its score
    return batch_detect_language_key_phrases_response.data


def classify_ai(data):
    # Classifies the main 5 key phrases
    data = str(data)
    json_dict = json.loads(data)
    dict4 = json_dict['documents'][0]['key_phrases'][1:5]
    # Return the 4 first main phrases from the dict (he passes the first word because we do not like her)
    return dict4


def google_search( value, api_key, cse_id, **kwargs):
    """this function took me soo mach time to make you
    need to create google api with Google cloud and your own search engine
    welcome to PANTHERSEARCH and no I didn't want to use selenium"""

    term = ""
    for word in value:
        term += str(word) + " "

    # the combine of the 4 main phrases for the search
    search_term = term
    service = build("customsearch", "v1", developerKey=api_key)
    req = service.cse().list(q=search_term, cx=cse_id, **kwargs).execute()
    # Returns the request == customsearch#search, the url we are searching
    return req['items']


def google_api_translate(fulltext):
    # translates the docx from hebrew to english
    to_translate = fulltext
    translated = GoogleTranslator(source='auto', target='en').translate(to_translate)
    return translated


def google_api_translate_revers(translated):
    to_translate = translated
    translated = GoogleTranslator(source='en', target='iw').translate(to_translate)
    print(translated)


def main():
    my_api_key = "AIzaSyAfAAwfwQFVCKNne9WJrNixONncLSgDp38"
    my_cse_id = "3028683e43d81493a"
    fulltext = read_docx()
    translated = google_api_translate(fulltext)
    config = create_config()
    data = api_ai(translated, config)
    key_phrases = Classify_ai(data)
    value = [prase.get('text') for prase in key_phrases]
    results = google_search(value, my_api_key, my_cse_id, num=10)
    # result is a line the Google gives us in the search
    for result in results:
        with open(TRUSTED_METADATA, "r") as trusted_f:
            tf = trusted_f.readlines()
            # line is a line in TRUSTED_METADATA file
            for line in tf:
                # Check if the current display link from the google search is equal to a line in the TRUSTED_METADATA
                if (result['displayLink'] in line):
                    print(result['title'], result['link'])
                else:
                    None


if __name__ == "__main__":
    main()
